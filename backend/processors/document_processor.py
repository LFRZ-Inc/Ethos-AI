"""
Document Processing System for Ethos AI
Handles PDFs, images, text files, and other document types
"""

import asyncio
import logging
import os
import tempfile
import json
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from pathlib import Path
import mimetypes
import hashlib
from datetime import datetime

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Metadata for processed documents"""
    filename: str
    file_type: str
    file_size: int
    pages: int
    word_count: int
    processing_time: float
    extracted_text: str
    summary: Optional[str] = None
    keywords: List[str] = None
    topics: List[str] = None
    language: str = 'en'
    confidence: float = 1.0
    error: Optional[str] = None

@dataclass
class KnowledgeEntry:
    """A single knowledge base entry"""
    id: str
    title: str
    content: str
    source: str
    created_at: datetime
    updated_at: datetime
    document_id: Optional[str] = None
    tags: List[str] = None
    metadata: Dict[str, Any] = None

class DocumentProcessor:
    """Main document processing system"""
    
    def __init__(self, orchestrator):
        self.orchestrator = orchestrator
        self.supported_formats = {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'txt': True,
            'md': True,
            'json': True,
            'csv': PANDAS_AVAILABLE,
            'jpg': OCR_AVAILABLE,
            'jpeg': OCR_AVAILABLE,
            'png': OCR_AVAILABLE,
            'gif': OCR_AVAILABLE,
            'bmp': OCR_AVAILABLE,
            'tiff': OCR_AVAILABLE
        }
        
        # Create temporary directory for processing
        self.temp_dir = Path(tempfile.mkdtemp(prefix="ethos_docs_"))
        logger.info(f"Document processor initialized with temp dir: {self.temp_dir}")
    
    async def process_document(self, file_path: str, file_content: bytes = None) -> DocumentMetadata:
        """Process a document and extract information"""
        try:
            start_time = datetime.now()
            
            # Determine file type
            file_type = self._get_file_type(file_path)
            if not self.supported_formats.get(file_type, False):
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Extract text based on file type
            if file_type == 'pdf':
                extracted_text = await self._extract_pdf_text(file_path, file_content)
            elif file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
                extracted_text = await self._extract_image_text(file_path, file_content)
            elif file_type == 'docx':
                extracted_text = await self._extract_docx_text(file_path, file_content)
            elif file_type == 'csv':
                extracted_text = await self._extract_csv_text(file_path, file_content)
            elif file_type in ['txt', 'md', 'json']:
                extracted_text = await self._extract_text_file(file_path, file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Analyze extracted text
            analysis = await self._analyze_text(extracted_text)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                file_type=file_type,
                file_size=len(file_content) if file_content else os.path.getsize(file_path),
                pages=analysis.get('pages', 1),
                word_count=analysis.get('word_count', 0),
                processing_time=processing_time,
                extracted_text=extracted_text,
                summary=analysis.get('summary'),
                keywords=analysis.get('keywords', []),
                topics=analysis.get('topics', []),
                language=analysis.get('language', 'en'),
                confidence=analysis.get('confidence', 1.0)
            )
            
            logger.info(f"Processed document: {metadata.filename} ({metadata.word_count} words)")
            return metadata
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return DocumentMetadata(
                filename=Path(file_path).name,
                file_type=self._get_file_type(file_path),
                file_size=0,
                pages=0,
                word_count=0,
                processing_time=0,
                extracted_text="",
                error=str(e)
            )
    
    def _get_file_type(self, file_path: str) -> str:
        """Get file type from extension"""
        return Path(file_path).suffix.lower().lstrip('.')
    
    async def _extract_pdf_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from PDF files"""
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing not available. Install PyPDF2 and pdfplumber")
        
        try:
            text_parts = []
            
            # Try pdfplumber first (better text extraction)
            try:
                if file_content:
                    with pdfplumber.open(file_content) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                else:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
                
                # Fallback to PyPDF2
                if file_content:
                    pdf_reader = PyPDF2.PdfReader(file_content)
                else:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    async def _extract_image_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from images using OCR"""
        if not OCR_AVAILABLE:
            raise ValueError("OCR not available. Install pytesseract and Pillow")
        
        try:
            if file_content:
                image = Image.open(file_content)
            else:
                image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting image text: {e}")
            raise
    
    async def _extract_docx_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Install python-docx")
        
        try:
            if file_content:
                doc = docx.Document(file_content)
            else:
                doc = docx.Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    async def _extract_csv_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from CSV files"""
        if not PANDAS_AVAILABLE:
            raise ValueError("CSV processing not available. Install pandas")
        
        try:
            if file_content:
                df = pd.read_csv(file_content)
            else:
                df = pd.read_csv(file_path)
            
            # Convert DataFrame to text representation
            text_parts = []
            
            # Add column names
            text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            text_parts.append(f"Rows: {len(df)}")
            text_parts.append("")
            
            # Add first few rows as sample
            sample_rows = df.head(10).to_string(index=False)
            text_parts.append("Sample data:")
            text_parts.append(sample_rows)
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("")
                text_parts.append("Summary statistics:")
                text_parts.append(df[numeric_cols].describe().to_string())
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting CSV text: {e}")
            raise
    
    async def _extract_text_file(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from plain text files"""
        try:
            if file_content:
                return file_content.decode('utf-8', errors='ignore')
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read()
        except Exception as e:
            logger.error(f"Error extracting text file: {e}")
            raise
    
    async def _analyze_text(self, text: str) -> Dict[str, Any]:
        """Analyze extracted text using AI"""
        try:
            if not text.strip():
                return {
                    'word_count': 0,
                    'pages': 0,
                    'summary': 'Empty document',
                    'keywords': [],
                    'topics': [],
                    'language': 'en',
                    'confidence': 0.0
                }
            
            # Use the orchestrator to analyze the text
            prompt = f"""
            Please analyze the following document text and provide:
            1. A concise summary (2-3 sentences)
            2. Key topics/themes discussed
            3. Important keywords
            4. Estimated number of pages (if applicable)
            5. Language detection
            6. Confidence score (0-1) for the analysis

            Document text:
            {text[:3000]}...

            Please respond in JSON format:
            {{
                "summary": "brief summary",
                "topics": ["topic1", "topic2"],
                "keywords": ["keyword1", "keyword2"],
                "pages": estimated_pages,
                "language": "en",
                "confidence": 0.95
            }}
            """
            
            response = await self.orchestrator.process_message(
                prompt,
                model_override='claude-3.5',  # Good for analysis
                use_tools=False
            )
            
            # Parse the response
            try:
                analysis = json.loads(response.content)
            except json.JSONDecodeError:
                # Fallback analysis
                analysis = {
                    'summary': 'Document analysis completed',
                    'topics': ['document'],
                    'keywords': text.split()[:10],
                    'pages': max(1, len(text) // 500),
                    'language': 'en',
                    'confidence': 0.8
                }
            
            # Add word count
            analysis['word_count'] = len(text.split())
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing text: {e}")
            return {
                'word_count': len(text.split()),
                'pages': max(1, len(text) // 500),
                'summary': 'Document analysis failed',
                'keywords': [],
                'topics': [],
                'language': 'en',
                'confidence': 0.5
            }

class KnowledgeBase:
    """Personal knowledge base management system"""
    
    def __init__(self, database):
        self.database = database
        self.entries: Dict[str, KnowledgeEntry] = {}
    
    async def add_entry(self, title: str, content: str, source: str, 
                       document_id: Optional[str] = None, tags: List[str] = None) -> KnowledgeEntry:
        """Add a new knowledge base entry"""
        try:
            entry_id = hashlib.md5(f"{title}{content}{source}".encode()).hexdigest()
            
            entry = KnowledgeEntry(
                id=entry_id,
                title=title,
                content=content,
                source=source,
                document_id=document_id,
                tags=tags or [],
                created_at=datetime.now(),
                updated_at=datetime.now(),
                metadata={}
            )
            
            self.entries[entry_id] = entry
            
            # Store in database
            await self._store_entry(entry)
            
            logger.info(f"Added knowledge entry: {title}")
            return entry
            
        except Exception as e:
            logger.error(f"Error adding knowledge entry: {e}")
            raise
    
    async def search_entries(self, query: str, limit: int = 10) -> List[KnowledgeEntry]:
        """Search knowledge base entries"""
        try:
            # Simple text search (can be enhanced with vector search)
            results = []
            query_lower = query.lower()
            
            for entry in self.entries.values():
                score = 0
                
                # Search in title
                if query_lower in entry.title.lower():
                    score += 3
                
                # Search in content
                if query_lower in entry.content.lower():
                    score += 1
                
                # Search in tags
                for tag in entry.tags:
                    if query_lower in tag.lower():
                        score += 2
                
                if score > 0:
                    results.append((entry, score))
            
            # Sort by score and return top results
            results.sort(key=lambda x: x[1], reverse=True)
            return [entry for entry, score in results[:limit]]
            
        except Exception as e:
            logger.error(f"Error searching knowledge base: {e}")
            return []
    
    async def get_entry(self, entry_id: str) -> Optional[KnowledgeEntry]:
        """Get a specific knowledge base entry"""
        return self.entries.get(entry_id)
    
    async def update_entry(self, entry_id: str, updates: Dict[str, Any]) -> bool:
        """Update a knowledge base entry"""
        try:
            if entry_id not in self.entries:
                return False
            
            entry = self.entries[entry_id]
            
            # Update fields
            for key, value in updates.items():
                if hasattr(entry, key):
                    setattr(entry, key, value)
            
            entry.updated_at = datetime.now()
            
            # Update in database
            await self._update_entry(entry)
            
            logger.info(f"Updated knowledge entry: {entry_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error updating knowledge entry: {e}")
            return False
    
    async def delete_entry(self, entry_id: str) -> bool:
        """Delete a knowledge base entry"""
        try:
            if entry_id not in self.entries:
                return False
            
            del self.entries[entry_id]
            
            # Remove from database
            await self._delete_entry(entry_id)
            
            logger.info(f"Deleted knowledge entry: {entry_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting knowledge entry: {e}")
            return False
    
    async def get_entries_by_tag(self, tag: str) -> List[KnowledgeEntry]:
        """Get all entries with a specific tag"""
        return [entry for entry in self.entries.values() if tag in entry.tags]
    
    async def get_all_tags(self) -> List[str]:
        """Get all unique tags"""
        tags = set()
        for entry in self.entries.values():
            tags.update(entry.tags)
        return sorted(list(tags))
    
    # Database operations
    async def _store_entry(self, entry: KnowledgeEntry) -> None:
        """Store entry in database"""
        # This would store the entry in the database
        # For now, we'll just log it
        logger.info(f"Storing knowledge entry: {entry.id}")
    
    async def _update_entry(self, entry: KnowledgeEntry) -> None:
        """Update entry in database"""
        # This would update the entry in the database
        # For now, we'll just log it
        logger.info(f"Updating knowledge entry: {entry.id}")
    
    async def _delete_entry(self, entry_id: str) -> None:
        """Delete entry from database"""
        # This would delete the entry from the database
        # For now, we'll just log it
        logger.info(f"Deleting knowledge entry: {entry_id}")

class CitationSystem:
    """Citation and source tracking system"""
    
    def __init__(self):
        self.citations: Dict[str, Dict[str, Any]] = {}
    
    def add_citation(self, source: str, content: str, metadata: Dict[str, Any] = None) -> str:
        """Add a new citation"""
        citation_id = hashlib.md5(f"{source}{content}".encode()).hexdigest()
        
        self.citations[citation_id] = {
            'source': source,
            'content': content,
            'metadata': metadata or {},
            'created_at': datetime.now().isoformat(),
            'usage_count': 0
        }
        
        return citation_id
    
    def get_citation(self, citation_id: str) -> Optional[Dict[str, Any]]:
        """Get citation by ID"""
        if citation_id in self.citations:
            self.citations[citation_id]['usage_count'] += 1
            return self.citations[citation_id]
        return None
    
    def search_citations(self, query: str) -> List[Dict[str, Any]]:
        """Search citations"""
        results = []
        query_lower = query.lower()
        
        for citation_id, citation in self.citations.items():
            if (query_lower in citation['source'].lower() or 
                query_lower in citation['content'].lower()):
                results.append({**citation, 'id': citation_id})
        
        return results 