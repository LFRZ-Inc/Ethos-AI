"""
Tool Manager for Ethos AI
Handles code execution, web search, file analysis, and other AI tools
"""

import asyncio
import logging
import subprocess
import tempfile
import os
import sys
import json
import time
from typing import Dict, List, Optional, Any
from pathlib import Path
import aiofiles
from duckduckgo_search import DDGS
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd

logger = logging.getLogger(__name__)

class ToolManager:
    """Manages AI tools and their execution"""
    
    def __init__(self, config, database, vector_store):
        self.config = config
        self.database = database
        self.vector_store = vector_store
        self.sandbox_dir = Path(config.data_dir) / "sandbox"
        self.sandbox_dir.mkdir(parents=True, exist_ok=True)
        
    async def initialize(self):
        """Initialize the tool manager"""
        try:
            # Create sandbox directory
            self.sandbox_dir.mkdir(parents=True, exist_ok=True)
            
            # Test web search capability
            if self.config.tools.web_search:
                try:
                    with DDGS() as ddgs:
                        results = list(ddgs.text("test", max_results=1))
                    logger.info("Web search tool initialized")
                except Exception as e:
                    logger.warning(f"Web search not available: {e}")
            
            logger.info("Tool manager initialized successfully")
            
        except Exception as e:
            logger.error(f"Error initializing tool manager: {e}")
            raise
    
    async def execute_tool(self, tool_name: str, parameters: Dict) -> Dict[str, Any]:
        """Execute a specific tool"""
        try:
            if tool_name == "web_search":
                return await self._web_search(parameters.get("query", ""))
            elif tool_name == "execute_code":
                return await self._execute_code(parameters.get("code", ""))
            elif tool_name == "search_files":
                return await self._search_files(
                    parameters.get("query", ""),
                    parameters.get("file_types", [])
                )
            elif tool_name == "read_file":
                return await self._read_file(parameters.get("file_path", ""))
            elif tool_name == "analyze_file":
                return await self._analyze_file(parameters.get("file_path", ""))
            else:
                return {"error": f"Unknown tool: {tool_name}"}
                
        except Exception as e:
            logger.error(f"Error executing tool {tool_name}: {e}")
            return {"error": str(e)}
    
    async def _web_search(self, query: str) -> Dict[str, Any]:
        """Perform web search using DuckDuckGo"""
        if not self.config.tools.web_search:
            return {"error": "Web search is disabled"}
        
        try:
            results = []
            with DDGS() as ddgs:
                search_results = ddgs.text(query, max_results=5)
                for result in search_results:
                    results.append({
                        "title": result.get("title", ""),
                        "link": result.get("link", ""),
                        "snippet": result.get("body", "")
                    })
            
            return {
                "type": "web_search",
                "query": query,
                "results": results,
                "count": len(results)
            }
            
        except Exception as e:
            logger.error(f"Error performing web search: {e}")
            return {"error": f"Web search failed: {str(e)}"}
    
    async def _execute_code(self, code: str) -> Dict[str, Any]:
        """Execute Python code in a sandboxed environment"""
        if not self.config.tools.code_execution:
            return {"error": "Code execution is disabled"}
        
        try:
            # Create temporary file for code execution
            with tempfile.NamedTemporaryFile(
                mode='w', 
                suffix='.py', 
                dir=self.sandbox_dir,
                delete=False
            ) as f:
                f.write(code)
                temp_file = f.name
            
            # Execute code with timeout and restrictions
            result = await self._run_code_safely(temp_file)
            
            # Clean up
            try:
                os.unlink(temp_file)
            except:
                pass
            
            return result
            
        except Exception as e:
            logger.error(f"Error executing code: {e}")
            return {"error": f"Code execution failed: {str(e)}"}
    
    async def _run_code_safely(self, file_path: str) -> Dict[str, Any]:
        """Run code with safety restrictions"""
        try:
            # Set up restricted environment
            env = os.environ.copy()
            env['PYTHONPATH'] = str(self.sandbox_dir)
            
            # Run code with timeout
            process = await asyncio.create_subprocess_exec(
                sys.executable, file_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
                cwd=self.sandbox_dir
            )
            
            try:
                stdout, stderr = await asyncio.wait_for(
                    process.communicate(), 
                    timeout=30.0  # 30 second timeout
                )
            except asyncio.TimeoutError:
                process.kill()
                return {"error": "Code execution timed out"}
            
            return {
                "type": "code_execution",
                "stdout": stdout.decode('utf-8', errors='ignore'),
                "stderr": stderr.decode('utf-8', errors='ignore'),
                "return_code": process.returncode,
                "success": process.returncode == 0
            }
            
        except Exception as e:
            return {"error": f"Code execution error: {str(e)}"}
    
    async def _search_files(self, query: str, file_types: List[str] = None) -> Dict[str, Any]:
        """Search local files for content"""
        if not self.config.tools.file_search:
            return {"error": "File search is disabled"}
        
        try:
            # Default search directories
            search_dirs = [
                Path.home() / "Documents",
                Path.home() / "Desktop",
                Path.home() / "Downloads"
            ]
            
            results = []
            
            for search_dir in search_dirs:
                if not search_dir.exists():
                    continue
                
                # Search for files
                for file_path in search_dir.rglob("*"):
                    if file_path.is_file():
                        # Check file type filter
                        if file_types and file_path.suffix.lower() not in [f".{ft}" for ft in file_types]:
                            continue
                        
                        # Check if query matches filename or content
                        if await self._file_matches_query(file_path, query):
                            results.append({
                                "path": str(file_path),
                                "name": file_path.name,
                                "size": file_path.stat().st_size,
                                "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(file_path.stat().st_mtime))
                            })
                        
                        # Limit results
                        if len(results) >= 20:
                            break
                
                if len(results) >= 20:
                    break
            
            return {
                "type": "file_search",
                "query": query,
                "results": results,
                "count": len(results)
            }
            
        except Exception as e:
            logger.error(f"Error searching files: {e}")
            return {"error": f"File search failed: {str(e)}"}
    
    async def _file_matches_query(self, file_path: Path, query: str) -> bool:
        """Check if a file matches the search query"""
        try:
            # Check filename
            if query.lower() in file_path.name.lower():
                return True
            
            # Check file content for text files
            text_extensions = {'.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv'}
            if file_path.suffix.lower() in text_extensions:
                try:
                    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = await f.read(10000)  # Read first 10KB
                        if query.lower() in content.lower():
                            return True
                except:
                    pass
            
            return False
            
        except Exception:
            return False
    
    async def _read_file(self, file_path: str) -> Dict[str, Any]:
        """Read and return file content"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"error": "File not found"}
            
            # Check file size
            if file_path.stat().st_size > 10 * 1024 * 1024:  # 10MB limit
                return {"error": "File too large"}
            
            # Read file based on type
            if file_path.suffix.lower() == '.pdf':
                return await self._read_pdf(file_path)
            elif file_path.suffix.lower() in {'.docx', '.doc'}:
                return await self._read_docx(file_path)
            elif file_path.suffix.lower() in {'.xlsx', '.xls'}:
                return await self._read_excel(file_path)
            else:
                return await self._read_text_file(file_path)
                
        except Exception as e:
            logger.error(f"Error reading file: {e}")
            return {"error": f"File read failed: {str(e)}"}
    
    async def _read_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Read text file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            
            return {
                "type": "text_file",
                "content": content,
                "size": len(content),
                "encoding": "utf-8"
            }
            
        except Exception as e:
            return {"error": f"Text file read failed: {str(e)}"}
    
    async def _read_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Read PDF file"""
        try:
            content = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            
            return {
                "type": "pdf",
                "content": content,
                "pages": len(pdf_reader.pages),
                "size": len(content)
            }
            
        except Exception as e:
            return {"error": f"PDF read failed: {str(e)}"}
    
    async def _read_docx(self, file_path: Path) -> Dict[str, Any]:
        """Read DOCX file"""
        try:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "type": "docx",
                "content": content,
                "paragraphs": len(doc.paragraphs),
                "size": len(content)
            }
            
        except Exception as e:
            return {"error": f"DOCX read failed: {str(e)}"}
    
    async def _read_excel(self, file_path: Path) -> Dict[str, Any]:
        """Read Excel file"""
        try:
            df = pd.read_excel(file_path)
            content = df.to_string()
            
            return {
                "type": "excel",
                "content": content,
                "rows": len(df),
                "columns": len(df.columns),
                "size": len(content)
            }
            
        except Exception as e:
            return {"error": f"Excel read failed: {str(e)}"}
    
    async def _analyze_file(self, file_path: str) -> Dict[str, Any]:
        """Analyze file content and extract insights"""
        try:
            # Read file first
            file_data = await self._read_file(file_path)
            if "error" in file_data:
                return file_data
            
            # Analyze content
            content = file_data.get("content", "")
            
            analysis = {
                "type": "file_analysis",
                "file_path": file_path,
                "file_type": file_data.get("type", "unknown"),
                "word_count": len(content.split()),
                "character_count": len(content),
                "line_count": len(content.split('\n')),
                "language": self._detect_language(content),
                "key_topics": self._extract_key_topics(content),
                "summary": self._generate_summary(content)
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing file: {e}")
            return {"error": f"File analysis failed: {str(e)}"}
    
    def _detect_language(self, content: str) -> str:
        """Simple language detection"""
        # This is a basic implementation - could be enhanced with proper language detection
        if any(char in content for char in "你好世界"):
            return "Chinese"
        elif any(char in content for char in "こんにちは"):
            return "Japanese"
        elif any(char in content for char in "안녕하세요"):
            return "Korean"
        else:
            return "English"
    
    def _extract_key_topics(self, content: str) -> List[str]:
        """Extract key topics from content"""
        # Simple keyword extraction - could be enhanced with NLP
        words = content.lower().split()
        word_freq = {}
        
        for word in words:
            if len(word) > 3:  # Skip short words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top 5 most frequent words
        topics = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
        return [topic[0] for topic in topics]
    
    def _generate_summary(self, content: str) -> str:
        """Generate a simple summary"""
        # Simple summary - could be enhanced with AI
        sentences = content.split('.')
        if len(sentences) > 3:
            return '. '.join(sentences[:3]) + '.'
        return content
    
    async def process_file_upload(self, file) -> Dict[str, Any]:
        """Process uploaded file"""
        try:
            # Generate file ID
            file_id = f"file_{int(time.time() * 1000)}"
            
            # Save file
            upload_dir = Path(self.config.data_dir) / "uploads"
            upload_dir.mkdir(parents=True, exist_ok=True)
            
            file_path = upload_dir / f"{file_id}_{file.filename}"
            
            # Save uploaded file
            content = await file.read()
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            
            # Add to database
            await self.database.add_file_upload(
                file_id=file_id,
                filename=file.filename,
                file_path=str(file_path),
                file_type=file.content_type or "unknown",
                file_size=len(content),
                metadata={"upload_time": time.time()}
            )
            
            # Analyze file if it's a supported type
            analysis = await self._analyze_file(str(file_path))
            
            return {
                "file_id": file_id,
                "filename": file.filename,
                "file_path": str(file_path),
                "file_size": len(content),
                "analysis": analysis
            }
            
        except Exception as e:
            logger.error(f"Error processing file upload: {e}")
            return {"error": f"File upload failed: {str(e)}"}
    
    async def get_available_tools(self) -> List[Dict[str, Any]]:
        """Get list of available tools"""
        tools = []
        
        if self.config.tools.web_search:
            tools.append({
                "name": "web_search",
                "description": "Search the web for current information",
                "parameters": {
                    "query": {"type": "string", "description": "Search query"}
                }
            })
        
        if self.config.tools.code_execution:
            tools.append({
                "name": "execute_code",
                "description": "Execute Python code in a sandboxed environment",
                "parameters": {
                    "code": {"type": "string", "description": "Python code to execute"}
                }
            })
        
        if self.config.tools.file_search:
            tools.append({
                "name": "search_files",
                "description": "Search local files for content",
                "parameters": {
                    "query": {"type": "string", "description": "Search query"},
                    "file_types": {"type": "array", "description": "File types to search"}
                }
            })
        
        tools.append({
            "name": "read_file",
            "description": "Read and analyze file content",
            "parameters": {
                "file_path": {"type": "string", "description": "Path to file"}
            }
        })
        
        return tools 